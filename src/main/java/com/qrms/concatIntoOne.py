from docx import Document
from os import walk
# walkDir = input('Enter dir name: ')
f = []
document = Document()

for (dirpath, dirnames, filenames) in walk('spring'):
	for f in filenames:
		with open(dirpath+'\\'+f) as contentFile:
			document.add_heading(dirpath+'\\'+f,level=1)
			document.add_paragraph(contentFile.read())

document.save('CPR.docx')